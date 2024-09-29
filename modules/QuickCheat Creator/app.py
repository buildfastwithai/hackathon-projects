import os
import streamlit as st
from langchain_google_genai import ChatGoogleGenerativeAI
from io import BytesIO
from fpdf import FPDF  
from docx import Document  
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
google_api_key = os.getenv("GOOGLE_API_KEY")

# CheatSheetGenerator class 
class CheatSheetGenerator:
    def __init__(self):
        # Initialize the LLM with the Google Gemini API
        self.llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            temperature=0.2,
            max_tokens=None,
            api_key=google_api_key
        )

    def generate_cheatsheet(self, topic):
        # Create a detailed prompt to generate an in-depth cheatsheet
        prompt = f"""
        Create a detailed and comprehensive last-minute preparation cheatsheet on '{topic}' with an extensive focus on the following aspects:
        - **Introduction**: Provide a thorough overview of the topic, including its significance and applications.
        - **Key Concepts**: Define and explain the fundamental concepts related to the topic. Include detailed explanations and examples for each concept.
        - **In-Depth Subtopics**: Break down the topic into multiple subtopics. For each subtopic:
          - Offer a comprehensive explanation.
          - Include relevant formulas, equations, or principles.
          - Discuss practical examples or case studies that illustrate the application of the subtopic.
        - **Common Challenges**: Highlight common misconceptions or challenges related to the topic and provide guidance on overcoming them.
        - **Interview Tips**: Offer strategic tips for discussing this topic in interviews, including potential questions and how to answer them effectively.
        - **FAQs**: Compile a list of frequently asked questions regarding the topic, with clear and informative answers.

        Ensure that the information is organized, structured, and presented in a clear manner. Use bullet points and numbered lists where applicable for easy reading.
        """
        
        messages = [("system", "You are a helpful assistant."), ("human", prompt)]
        response = self.llm.invoke(messages).content
        return response

# DocumentHandler class to handle export functionality
class DocumentHandler:
    def export_to_docx(self, cheatsheet_text, filename):
        # Create a Word document with the generated cheatsheet text
        doc = Document()
        doc.add_heading('Cheat Sheet', 0)
        doc.add_paragraph(cheatsheet_text)

        # Save the document to a BytesIO object for downloading
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def export_to_pdf(self, cheatsheet_text, filename):
        # Create a PDF document with the cheatsheet text
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, cheatsheet_text)

        # Save the PDF content to a BytesIO object for downloading
        buffer = BytesIO()
        pdf_output = pdf.output(dest='S').encode('latin1')  # Get the content as a byte string
        buffer.write(pdf_output)
        buffer.seek(0)  # Move the cursor to the start of the buffer
        return buffer

# Streamlit application layout
def main():
    st.set_page_config(page_title="QuickCheat Creator", layout="wide")  # Set page title and layout

    st.title("📝 QuickCheat Creator")
    st.sidebar.header("Generate Your Instant Cheat Sheet")  # Sidebar header

    st.markdown(""" 
        Supercharge your exam preparation with **QuickCheat Creator**, your go-to tool for generating instant cheat sheets! 
        Simply enter the topic name, and watch as our intelligent AI crafts a concise and organized summary of key concepts and essential information tailored to that subject. 
        No need for extensive study materials—just a few clicks, and you’ll have a personalized cheat sheet at your fingertips. 
        Perfect for quick reviews and last-minute studying, **QuickCheat Creator** helps you maximize your study efficiency and approach your exams with confidence!
    """)

    # Input field for topic in the sidebar
    topic = st.sidebar.text_input("Enter Topic:", placeholder="e.g., Machine Learning")
    
    # Button to generate cheat sheet in the sidebar
    if st.sidebar.button("Generate Cheat Sheet"):
        if topic:
            with st.spinner("Generating cheat sheet..."):
                generator = CheatSheetGenerator()
                cheat_sheet = generator.generate_cheatsheet(topic)

            # Store the cheat sheet in session state to prevent it from disappearing
            st.session_state.cheat_sheet = cheat_sheet

    # Display the cheat sheet if it exists
    if 'cheat_sheet' in st.session_state:
        st.subheader("Generated Cheat Sheet")
        st.write(st.session_state.cheat_sheet)

        # Create DocumentHandler instance
        doc_handler = DocumentHandler()

        # PDF download option in the sidebar
        pdf_file = doc_handler.export_to_pdf(st.session_state.cheat_sheet, filename=f"{topic}.pdf")
        st.sidebar.download_button("Download as PDF", pdf_file, f"{topic}.pdf", key="pdf_download")

        # DOCX download option in the sidebar
        docx_file = doc_handler.export_to_docx(st.session_state.cheat_sheet, filename=f"{topic}.docx")
        st.sidebar.download_button("Download as DOCX", docx_file, f"{topic}.docx", key="docx_download")

    else:
        st.warning("Please enter a topic and generate the cheat sheet.")


if __name__ == "__main__":
    main()
